"""
Word document building-blocks.
Every helper takes a Document (or paragraph) and mutates it in-place.
word_report.py imports everything from here — no docx code lives elsewhere.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ── Colour palette ────────────────────────────────────────────────────────────
COL_HEADING = RGBColor(0x1F, 0x39, 0x64)   # dark navy
COL_SUBHEAD = RGBColor(0x2E, 0x74, 0xB5)   # medium blue
COL_BODY    = RGBColor(0x00, 0x00, 0x00)   # black
COL_CAPTION = RGBColor(0x40, 0x40, 0x40)   # dark grey
COL_MISS    = RGBColor(0x80, 0x80, 0x80)   # placeholder grey
COL_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
HDR_BG_HEX  = "1F3964"                     # table header background (hex str)


# ── Document setup ────────────────────────────────────────────────────────────

def set_doc_margins(doc: Document) -> None:
    """A4 page, 2.5 cm margins."""
    sec = doc.sections[0]
    sec.page_height   = Cm(29.7)
    sec.page_width    = Cm(21.0)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)


def set_default_font(doc: Document, name: str = "Calibri", size: int = 10) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = name
    normal.font.size = Pt(size)


def add_header_footer(doc: Document, text: str) -> None:
    sec = doc.sections[0]
    header = sec.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(para, text, size=8, color=RGBColor(0x60, 0x60, 0x60), italic=True)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


# ── Text helpers ──────────────────────────────────────────────────────────────

def add_run(
    para,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size: int = 10,
    color: RGBColor | None = None,
    font_name: str = "Calibri",
    dynamic: bool = False,
):
    """
    Add a formatted run to an existing paragraph.
    dynamic=True → bold blue (used for site-specific values, matching original style).
    """
    run = para.add_run(text)
    run.italic     = italic
    run.font.size  = Pt(size)
    run.font.name  = font_name
    if dynamic:
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)
    else:
        run.bold = bold
        if color:
            run.font.color.rgb = color
    return run


def add_heading(
    doc: Document,
    text: str,
    size: int = 13,
    color: RGBColor | None = None,
    space_before: int = 12,
    space_after: int = 6,
):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    add_run(para, text, bold=True, size=size, color=color or COL_HEADING)
    return para


def add_subheading(
    doc: Document,
    text: str,
    size: int = 11,
    color: RGBColor | None = None,
    space_before: int = 8,
    space_after: int = 4,
):
    return add_heading(
        doc, text, size=size,
        color=color or COL_SUBHEAD,
        space_before=space_before,
        space_after=space_after,
    )


def add_body(
    doc: Document,
    text: str,
    size: int = 10,
    space_before: int = 2,
    space_after: int = 4,
    italic: bool = False,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    para.alignment = align
    add_run(para, text, size=size, italic=italic)
    return para


def add_bullet(doc: Document, text: str, size: int = 10) -> None:
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    add_run(para, text, size=size)


def add_caption(doc: Document, text: str) -> None:
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(cp, text, italic=True, size=9, color=COL_CAPTION)


# ── Image helper ──────────────────────────────────────────────────────────────

def add_image(
    doc: Document,
    img_path: Path,
    width=Inches(6),
    caption: str | None = None,
) -> None:
    img_path = Path(img_path)
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f"[Figure not available: {img_path.name}]", italic=True, size=9, color=COL_MISS)
    if caption:
        add_caption(doc, caption)


# ── Table helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background via low-level XML (python-docx limitation workaround)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def style_cell(
    cell,
    text: str,
    bold: bool = False,
    size: int = 9,
    bg_color: str | None = None,
    font_color: RGBColor | None = None,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    italic: bool = False,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    para.clear()
    add_run(para, text, bold=bold, size=size, color=font_color or COL_BODY, italic=italic)
    if bg_color:
        set_cell_bg(cell, bg_color)


# ── 2×2 chart grid ────────────────────────────────────────────────────────────

def insert_2x2_chart_grid(
    doc: Document,
    charts: list[tuple[Path, str]],
) -> None:
    """
    Insert a 2×2 table of chart images with captions.
    charts: list of exactly 4 (image_path, caption_text) tuples.
    Missing images are replaced with a placeholder label.
    """
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    positions = [(0, 0), (0, 1), (1, 0), (1, 1)]
    for (row_i, col_i), (img_path, caption) in zip(positions, charts[:4]):
        cell = table.cell(row_i, col_i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_path = Path(img_path)
        if img_path.exists():
            p.add_run().add_picture(str(img_path), width=Inches(3.0))
        else:
            add_run(p, f"[{img_path.name}]", italic=True, size=8, color=COL_MISS)
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cp, caption, italic=True, size=8, color=COL_CAPTION)
        cell.add_paragraph()   # bottom padding
