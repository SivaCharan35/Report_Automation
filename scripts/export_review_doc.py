"""
Offline review export — outside the main pipeline.

Merges all ara_*.json section outputs from a Report_* folder into:
  - final_report.json   (combined TipTap doc)
  - final_report.docx   (viewable Word preview for local QA)

Usage
─────
    python scripts/export_review_doc.py
    python scripts/export_review_doc.py Report_Data/Report_20260715_170021
    python scripts/export_review_doc.py --latest

Notes
─────
- Prefer local map paths under assets/ (often stored in image attrs.title
  when running --local-only).
- This is a QA renderer, not a pixel-perfect FE clone. Good enough to check
  text, tables, maps, and placeholders before BE/FE catch up.
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import sys
from pathlib import Path

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("export_review_doc")

# Repo root = parent of scripts/
_REPO_ROOT = Path(__file__).resolve().parent.parent
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

import config  # noqa: E402


# ── Load / merge ──────────────────────────────────────────────────────────────

def _latest_report_dir() -> Path:
    root = config.REPORT_DATA_DIR
    dirs = sorted(
        (p for p in root.glob("Report_*") if p.is_dir()),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not dirs:
        raise FileNotFoundError(f"No Report_* folders under {root}")
    return dirs[0]


def _load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def _ordered_sections(report_dir: Path) -> list[tuple[str, str, Path]]:
    """
    Return [(module, display_name, json_path), ...] in sequence order.
    Prefer _section_metadata.json; fall back to sorted ara_*.json.
    """
    meta_path = report_dir / "_section_metadata.json"
    if meta_path.exists():
        meta = _load_json(meta_path)
        rows = []
        for module, info in meta.items():
            path = report_dir / f"{module}.json"
            if not path.exists():
                logger.warning("Missing section file for %s — skipped", module)
                continue
            seq = info.get("sequence_number", 999)
            name = info.get("section_title") or info.get("name") or module
            rows.append((seq, module, name, path))
        rows.sort(key=lambda r: (r[0], r[1]))
        return [(m, n, p) for _, m, n, p in rows]

    paths = sorted(report_dir.glob("ara_*.json"))
    return [(p.stem, p.stem, p) for p in paths]


def merge_sections(report_dir: Path) -> dict:
    """Build one TipTap doc: {type: doc, content: [...all sections...]}."""
    sections = _ordered_sections(report_dir)
    merged_content: list = []
    section_index: list[dict] = []

    for module, name, path in sections:
        doc = _load_json(path)
        if isinstance(doc, str):
            try:
                doc = json.loads(doc)
            except json.JSONDecodeError:
                logger.warning("Unparseable string payload in %s", path.name)
                continue

        # Normal TipTap root
        if isinstance(doc, dict) and doc.get("type") == "doc":
            nodes = doc.get("content") or []
        elif isinstance(doc, dict) and "content" in doc:
            nodes = doc.get("content") or []
        elif isinstance(doc, dict) and "resolved_content" in doc:
            try:
                inner = json.loads(doc["resolved_content"])
                nodes = inner.get("content") or []
            except (json.JSONDecodeError, TypeError, AttributeError):
                logger.warning("Bad resolved_content in %s", path.name)
                continue
        else:
            logger.warning("Unrecognized shape in %s — skipped", path.name)
            continue

        section_index.append({
            "module": module,
            "name": name,
            "node_count": len(nodes),
            "source": path.name,
        })
        merged_content.extend(nodes)

    input_config = {}
    ic_path = report_dir / "_input_config.json"
    if ic_path.exists():
        input_config = _load_json(ic_path)

    return {
        "type": "doc",
        "meta": {
            "report_dir": str(report_dir),
            "input_config": input_config,
            "sections": section_index,
        },
        "content": merged_content,
    }


# ── TipTap → DOCX ─────────────────────────────────────────────────────────────

def _collect_text(node) -> str:
    if not isinstance(node, dict):
        return ""
    if node.get("type") == "text":
        return node.get("text") or ""
    parts = []
    for child in node.get("content") or []:
        parts.append(_collect_text(child))
    return "".join(parts)


def _heading_level(node: dict) -> int | None:
    attrs = node.get("attrs") or {}
    style = (attrs.get("styleId") or "").lower()
    if style == "heading1":
        return 1
    if style == "heading2":
        return 2
    if style == "heading3":
        return 3
    if style == "heading4":
        return 4
    if node.get("type") == "heading":
        try:
            return int(attrs.get("level") or 1)
        except (TypeError, ValueError):
            return 1
    return None


def _resolve_image_path(attrs: dict, report_dir: Path) -> Path | None:
    """
    Prefer local paths written by --local-only runs.
    Pipeline often puts the local PNG in attrs.title while attrs.src is a
    leftover Azure/BE URL.
    """
    candidates = [
        attrs.get("title"),
        attrs.get("src"),
        attrs.get("originalSrc"),
    ]
    for raw in candidates:
        if not isinstance(raw, str) or not raw.strip():
            continue
        s = raw.strip()
        # Local absolute / relative path
        p = Path(s)
        if p.exists() and p.is_file():
            return p
        # Filename only → look under report assets/
        name = Path(s).name
        if name.lower().endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
            for folder in (report_dir / "assets", report_dir):
                cand = folder / name
                if cand.exists():
                    return cand
    return None


def _add_runs(paragraph, node: dict) -> None:
    """Flatten tipTap text/run tree into a python-docx paragraph."""
    from docx.shared import Pt, RGBColor

    def walk(n, marks=None):
        marks = marks or {}
        if not isinstance(n, dict):
            return
        t = n.get("type")
        if t == "text":
            text = n.get("text") or ""
            if not text:
                return
            run = paragraph.add_run(text)
            run.bold = bool(marks.get("bold"))
            run.italic = bool(marks.get("italic"))
            # pick up tipTap textStyle marks if present on this node
            for m in n.get("marks") or []:
                if m.get("type") == "bold":
                    run.bold = True
                elif m.get("type") == "italic":
                    run.italic = True
                elif m.get("type") == "textStyle":
                    a = m.get("attrs") or {}
                    if a.get("fontSize"):
                        try:
                            sz = float(str(a["fontSize"]).replace("pt", ""))
                            run.font.size = Pt(sz)
                        except ValueError:
                            pass
                    if a.get("color"):
                        hx = str(a["color"]).lstrip("#")
                        if len(hx) == 6:
                            run.font.color.rgb = RGBColor(
                                int(hx[0:2], 16), int(hx[2:4], 16), int(hx[4:6], 16)
                            )
            return

        # Accumulate marks from nested "run" wrappers
        child_marks = dict(marks)
        if t == "run":
            for m in n.get("marks") or []:
                if m.get("type") == "bold":
                    child_marks["bold"] = True
                elif m.get("type") == "italic":
                    child_marks["italic"] = True

        # Inline images inside paragraph
        if t == "image":
            return  # handled at paragraph level separately

        for c in n.get("content") or []:
            walk(c, child_marks)

    walk(node)


def _paragraph_has_image(node: dict) -> list[dict]:
    imgs = []

    def walk(n):
        if isinstance(n, dict):
            if n.get("type") == "image":
                imgs.append(n)
            for c in n.get("content") or []:
                walk(c)

    walk(node)
    return imgs


def _render_table(doc, table_node: dict) -> None:
    rows = table_node.get("content") or []
    if not rows:
        return
    # Determine max cols
    n_cols = 0
    grid = []
    for row in rows:
        cells = row.get("content") or []
        n_cols = max(n_cols, len(cells))
        grid.append([_collect_text(c).strip() for c in cells])
    if n_cols == 0:
        return
    table = doc.add_table(rows=len(grid), cols=n_cols)
    table.style = "Table Grid"
    for r_i, row_vals in enumerate(grid):
        for c_i in range(n_cols):
            cell = table.cell(r_i, c_i)
            cell.text = row_vals[c_i] if c_i < len(row_vals) else ""
            if r_i == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def render_docx(merged: dict, report_dir: Path, out_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    from core.word_utils import add_header_footer, set_default_font, set_doc_margins

    doc = Document()
    set_doc_margins(doc)
    set_default_font(doc)
    area = (merged.get("meta") or {}).get("input_config", {}).get("area") or report_dir.name
    add_header_footer(doc, f"Local QA preview — {area}")

    title = doc.add_paragraph()
    run = title.add_run(f"Local Review Document — {area}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)

    note = doc.add_paragraph()
    nr = note.add_run(
        "Generated offline from merged ara_*.json. For internal validation only "
        "(not a FE pixel clone)."
    )
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for node in merged.get("content") or []:
        if not isinstance(node, dict):
            continue
        ntype = node.get("type")

        if ntype == "table":
            _render_table(doc, node)
            doc.add_paragraph("")
            continue

        if ntype == "bulletList":
            for item in node.get("content") or []:
                text = _collect_text(item).strip()
                if text:
                    doc.add_paragraph(text, style="List Bullet")
            continue

        if ntype == "orderedList":
            for item in node.get("content") or []:
                text = _collect_text(item).strip()
                if text:
                    doc.add_paragraph(text, style="List Number")
            continue

        if ntype in ("paragraph", "heading"):
            text = _collect_text(node).strip()
            imgs = _paragraph_has_image(node)
            level = _heading_level(node)

            # Image-only or image+caption paragraph
            if imgs:
                for img in imgs:
                    attrs = img.get("attrs") or {}
                    local = _resolve_image_path(attrs, report_dir)
                    if local is None:
                        miss = doc.add_paragraph()
                        mr = miss.add_run(
                            f"[image missing — title={attrs.get('title')!r} src={attrs.get('src')!r}]"
                        )
                        mr.italic = True
                        mr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                        continue
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run()
                        run.add_picture(str(local), width=Inches(5.8))
                        logger.info("Embedded image: %s", local.name)
                    except Exception as exc:
                        logger.warning("Could not embed %s: %s", local, exc)
                        p = doc.add_paragraph(f"[failed to embed {local.name}: {exc}]")

            if text:
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                elif level == 3:
                    doc.add_heading(text, level=3)
                elif level == 4:
                    doc.add_heading(text, level=4)
                else:
                    para = doc.add_paragraph()
                    _add_runs(para, node)
            continue

        # Fallback: dump visible text if any
        text = _collect_text(node).strip()
        if text:
            doc.add_paragraph(text)

    # Unresolved placeholder scan appendix
    blob = json.dumps(merged)
    leftovers = sorted(set(re.findall(r"\{\{[A-Z0-9_]+\}\}", blob)))
    doc.add_page_break()
    doc.add_heading("QA checklist", level=1)
    doc.add_paragraph(f"Report folder: {report_dir}")
    doc.add_paragraph(f"Sections merged: {len((merged.get('meta') or {}).get('sections') or [])}")
    if leftovers:
        doc.add_paragraph("Unresolved placeholders still present:")
        for tok in leftovers:
            doc.add_paragraph(tok, style="List Bullet")
    else:
        doc.add_paragraph("No unresolved {{PLACEHOLDER}} tokens found in merged JSON.")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    logger.info("Wrote %s", out_path)


# ── CLI ───────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(description="Merge report JSONs and export a review DOCX")
    parser.add_argument(
        "report_dir",
        nargs="?",
        default=None,
        help="Path to Report_YYYYMMDD_HHMMSS (default: latest under Report_Data/)",
    )
    parser.add_argument(
        "--latest",
        action="store_true",
        help="Explicitly use the newest Report_* folder",
    )
    parser.add_argument(
        "--json-only",
        action="store_true",
        help="Write final_report.json only (skip DOCX)",
    )
    args = parser.parse_args()

    if args.report_dir:
        report_dir = Path(args.report_dir)
        if not report_dir.is_absolute():
            report_dir = (_REPO_ROOT / report_dir).resolve()
    else:
        report_dir = _latest_report_dir()

    if not report_dir.exists():
        logger.error("Report dir not found: %s", report_dir)
        sys.exit(1)

    logger.info("Merging sections from %s", report_dir)
    merged = merge_sections(report_dir)

    json_out = report_dir / "final_report.json"
    json_out.write_text(json.dumps(merged, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info(
        "Wrote %s (%d content nodes, %d sections)",
        json_out.name,
        len(merged.get("content") or []),
        len((merged.get("meta") or {}).get("sections") or []),
    )

    if args.json_only:
        return

    docx_out = report_dir / "final_report.docx"
    render_docx(merged, report_dir, docx_out)
    print(f"\nReview files ready:\n  {json_out}\n  {docx_out}\n")


if __name__ == "__main__":
    main()
