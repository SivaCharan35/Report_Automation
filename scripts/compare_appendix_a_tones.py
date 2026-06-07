"""
Module: compare_appendix_a_tones

Standalone validation tool: generates a BEFORE vs AFTER comparison of
Appendix A impact text for every data layer, by calling Claude twice per
layer — once with the OLD (pre-rewrite, technical) prompt and once with
the NEW plain-English prompt currently in scripts/ara_risk_insights.py.

It does NOT run the full 13-section pipeline. It only exercises the
Appendix A impact-text generation logic. Outputs a Word document to:

    Report_Data/Comparison_<area>_<timestamp>.docx

with one page per layer showing the two Claude outputs side-by-side.

Usage
─────
    # Location 1 — defaults (uses current COG/Flood + COG/Heat folders):
    python -m scripts.compare_appendix_a_tones

    # Location 2 — swap COG/ files first, then re-run with new metadata:
    python -m scripts.compare_appendix_a_tones \
        --area "Shell Norco" --city "Norco" --state "Louisiana" --country "USA"

Workflow
────────
1. Run once with defaults → produces Comparison_Alagiyanallur_<ts>.docx
2. Swap COG/Flood, COG/Heat, COG/roads, COG/waterline files for the new area
3. Run again with --area / --city / --state / --country flags
4. Open both .docx files side-by-side and validate the prompt rewrite.

Cost note: ~18 Claude calls per location (9 layers × 2 prompts). Roughly
$0.20–$0.50 per location at current Claude Opus pricing.
"""

from __future__ import annotations

import argparse
import json
import logging
import sys
from datetime import datetime
from pathlib import Path

# Allow `python scripts/compare_appendix_a_tones.py` from project root.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import numpy as np
import rasterio

import config
from core.classifiers import classify_current_flood, classify_current_heat
from core.geojson_utils import compute_risk_counts, load_geojson
from scripts.ara_risk_insights import (
    _GEOJSON_CATALOGUE,
    _LAYER_CATALOGUE,
    _LAYER_EXTRA_CONTEXT,
    _LAYER_METRIC_GUIDE,
    _LST_BINS,
    _ask_claude_raster,
    _ask_claude_vector,
    _compute_road_stats,
    _compute_stats,
    _compute_water_stats,
    _find_geojson,
    _find_tif,
    _overall_hazard_profile,
    _severity_signal,
    _strip_markdown_emphasis,
    _susc_bins,
    _susc_lulc,
    _susc_percentile,
)

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
# Frozen OLD prompts (pre-rewrite — formal/technical voice)
# ─────────────────────────────────────────────────────────────────────────────

_OLD_RASTER_PROMPT = """You are a senior climate risk analyst writing the Appendix A section of a formal Asset Resilience Assessment report.

SITE: {site_name}
LOCATION: {location}
HAZARD: {hazard}
DATA LAYER: {layer_name}
RISK SIGNAL: {severity} — High-Susceptibility class covers {high_pct:.1f}% of the site
DOMINANT CLASS: {dominant_class} ({dominant_pct:.1f}% of site area)

--- RASTER STATISTICS ---
{stats_json}

--- SUSCEPTIBILITY CLASS DISTRIBUTION ---
{susc_json}

--- LAYER METRIC GUIDE ---
{metric_guide}

Write a paragraph of exactly 3–4 sentences (60–90 words).
Sentence 1: State the dominant susceptibility class and its exact % — open with the {severity} risk signal.
Sentence 2: Cite a specific metric and state its direct implication.
Sentence 3: Explain the physical mechanism driving {hazard_lower} risk at {site_name}.
{sentence_4_rule}

Rules: professional formal tone, every number cited must come from the data above, no bullet points, output only the paragraph."""


_OLD_VECTOR_PROMPT = """You are a senior climate risk analyst writing Appendix A of a formal Asset Resilience Assessment.

SITE: {site_name}
LOCATION: {location}
HAZARD: {hazard}
DATA LAYER: {layer_name}
DOMINANT FEATURE TYPE: {dom_type} ({dom_pct:.1f}% of {total_feat} total features)

--- FEATURE STATISTICS ---
{type_stats_json}

--- LAYER METRIC GUIDE ---
{metric_guide}

Write a paragraph of exactly 3–4 sentences (60–90 words).
Sentence 1: State the dominant feature type and its share.
Sentence 2: Cite a specific metric and state its direct flood risk implication.
Sentence 3: Explain the physical mechanism driving {hazard_lower} risk at {site_name}.
Sentence 4 (optional): State the primary asset resilience implication if risk is significant.

Rules: professional formal tone, every number cited must come from the data above, output only the paragraph."""


def _detect_geojson(directory: Path, keyword: str) -> Path | None:
    """Find the first *.geojson in `directory` whose filename contains `keyword`."""
    if not directory.exists():
        return None
    for f in sorted(directory.glob("*.geojson")):
        if keyword.lower() in f.name.lower():
            return f
    return None


def _inject_overall_risk_profiles(site_info: dict, input_files_dir: Path) -> None:
    """Load flood + heat building-level GeoJSONs from Input_Files, derive each
    hazard's overall profile + severity (via `_overall_hazard_profile`), and
    stuff the results into `site_info` so the AFTER prompts pick them up.

    Falls back to the neutral "no data available, moderate" stub when the
    Input_Files don't exist for this area — the tool still runs end-to-end."""
    flood_path = _detect_geojson(input_files_dir, "flood")
    heat_path  = _detect_geojson(input_files_dir, "heat")

    flood_counts: dict = {}
    heat_counts:  dict = {}

    if flood_path is not None:
        flood_data = load_geojson(flood_path)
        if flood_data:
            try:
                flood_counts = compute_risk_counts(flood_data, classify_current_flood)
            except Exception as exc:
                logger.warning("Could not derive flood risk counts: %s", exc)
    else:
        logger.warning("No flood GeoJSON in %s — overall flood profile defaults to neutral.", input_files_dir)

    if heat_path is not None:
        heat_data = load_geojson(heat_path)
        if heat_data:
            try:
                heat_counts = compute_risk_counts(heat_data, classify_current_heat)
            except Exception as exc:
                logger.warning("Could not derive heat risk counts: %s", exc)
    else:
        logger.warning("No heat GeoJSON in %s — overall heat profile defaults to neutral.", input_files_dir)

    site_info["flood_profile"], site_info["flood_overall"] = _overall_hazard_profile(flood_counts)
    site_info["heat_profile"],  site_info["heat_overall"]  = _overall_hazard_profile(heat_counts)


def _call_claude_with_old_prompt(prompt: str) -> str | None:
    """Direct Claude call for the OLD-prompt path. Same model + max_tokens as
    the live function. Markdown emphasis (e.g. `**ELEVATED**`) is stripped so
    the BEFORE column in the comparison doc reads consistently with the AFTER
    column, even though the OLD prompt didn't explicitly forbid Markdown."""
    try:
        import anthropic
        client = anthropic.Anthropic()
        msg = client.messages.create(
            model="claude-opus-4-6",
            max_tokens=200,
            messages=[{"role": "user", "content": prompt}],
        )
        return _strip_markdown_emphasis(msg.content[0].text.strip())
    except Exception as exc:
        logger.warning("    OLD-prompt Claude call failed: %s", exc)
        return None


# ─────────────────────────────────────────────────────────────────────────────
# Per-layer processing
# ─────────────────────────────────────────────────────────────────────────────

def _read_raster(tif_path: Path):
    try:
        with rasterio.open(tif_path) as src:
            data = src.read(1).astype(np.float64)
            nd = src.nodata
        if nd is not None:
            data[data == nd] = np.nan
        data[data <= -9000] = np.nan
        return data
    except Exception as exc:
        logger.error("    Cannot read %s: %s", tif_path.name, exc)
        return None


def _process_raster_layer(
    key: str, hazard: str, display_name: str, mode,
    cog_dir: Path, site_info: dict,
) -> dict:
    out: dict = {
        "layer_key":       key,
        "layer_name":      display_name,
        "hazard":          hazard,
        "before":          None,
        "after":           None,
        "severity":        None,
        "dominant":        None,
        "skipped_reason":  None,
    }

    tif = _find_tif(cog_dir / hazard, key)
    if tif is None:
        out["skipped_reason"] = f"no TIF for '{key}' in {cog_dir / hazard}"
        logger.warning("  ⚠ %s", out["skipped_reason"])
        return out

    data = _read_raster(tif)
    if data is None:
        out["skipped_reason"] = f"could not read {tif.name}"
        return out

    if mode == "lst_auto":
        valid = data[~np.isnan(data)]
        if valid.size and np.nanmean(valid) > 100:
            data -= 273.15  # Kelvin → Celsius

    try:
        if mode == "categorical":
            susc = _susc_lulc(data)
        elif isinstance(mode, list):
            susc = _susc_bins(data, mode)
        elif mode == "lst_auto":
            susc = _susc_bins(data, _LST_BINS)
        elif mode == "percentile_inverse":
            susc = _susc_percentile(data, "inverse")
        elif mode == "percentile_normal":
            susc = _susc_percentile(data, "normal")
        else:
            susc = {}
    except Exception as exc:
        out["skipped_reason"] = f"susceptibility computation failed: {exc}"
        return out

    stats = _compute_stats(data)
    severity, high_pct, dom_class, dom_pct = _severity_signal(key, susc, hazard)
    out["severity"] = severity
    out["dominant"] = f"{dom_class} ({dom_pct:.1f}%)"

    # ── BEFORE: OLD prompt ───────────────────────────────────────────────────
    guide_by_hazard = _LAYER_METRIC_GUIDE.get(key, {})
    metric_guide = (
        guide_by_hazard.get(hazard)
        or next(iter(guide_by_hazard.values()), None)
        or _LAYER_EXTRA_CONTEXT.get(key, "")
    )
    sentence_4_rule_old = (
        "Sentence 4: State the primary asset resilience implication or intervention priority."
        if severity in ("CRITICAL", "ELEVATED")
        else "Sentence 4: Omit — stop after Sentence 3."
    )
    city = site_info.get("city", "")
    country = site_info.get("country", "")
    location = ", ".join(p for p in [city, country] if p) or "the assessed location"

    old_prompt = _OLD_RASTER_PROMPT.format(
        site_name=site_info.get("site_name", "the assessed site"),
        location=location,
        hazard=hazard,
        hazard_lower=hazard.lower(),
        layer_name=display_name,
        severity=severity,
        high_pct=high_pct,
        dominant_class=dom_class,
        dominant_pct=dom_pct,
        stats_json=json.dumps(stats, indent=2),
        susc_json=json.dumps(susc, indent=2),
        metric_guide=metric_guide,
        sentence_4_rule=sentence_4_rule_old,
    )
    logger.info("    Calling Claude with OLD prompt...")
    out["before"] = _call_claude_with_old_prompt(old_prompt) or "(OLD Claude call failed)"

    # ── AFTER: NEW prompt via the live ara_risk_insights function ───────────
    logger.info("    Calling Claude with NEW prompt...")
    out["after"] = (
        _ask_claude_raster(display_name, hazard, stats, susc, site_info, key)
        or "(NEW Claude call failed)"
    )
    return out


def _process_vector_layer(
    key: str, subdir_name: str, display_name: str,
    cog_dir: Path, site_info: dict,
) -> dict:
    out: dict = {
        "layer_key":      key,
        "layer_name":     display_name,
        "hazard":         "Flood",
        "before":         None,
        "after":          None,
        "severity":       None,
        "dominant":       None,
        "skipped_reason": None,
    }

    try:
        import geopandas as gpd
    except ImportError:
        out["skipped_reason"] = "geopandas not installed"
        return out

    geojson_path = _find_geojson(cog_dir / subdir_name)
    if geojson_path is None:
        out["skipped_reason"] = f"no GeoJSON in {cog_dir / subdir_name}"
        logger.warning("  ⚠ %s", out["skipped_reason"])
        return out

    try:
        gdf = gpd.read_file(geojson_path)
    except Exception as exc:
        out["skipped_reason"] = f"cannot read {geojson_path.name}: {exc}"
        return out

    if gdf.empty:
        out["skipped_reason"] = "empty geojson"
        return out

    if gdf.crs is None:
        gdf = gdf.set_crs(4326)
    elif gdf.crs.to_epsg() != 4326:
        gdf = gdf.to_crs(4326)

    if key == "roads":
        type_stats = _compute_road_stats(gdf)
        hazard_label = "Flood and Heat"
    else:
        type_stats = _compute_water_stats(gdf)
        hazard_label = "Flood"
    out["hazard"] = hazard_label

    by_type = type_stats.get("by_type", {})
    if by_type:
        dom_type = max(by_type, key=lambda k: by_type[k].get("count", 0))
        dom_pct = by_type[dom_type].get("pct", 0.0)
    else:
        dom_type = "N/A"
        dom_pct = 0.0
    total_feat = type_stats.get("total_features", 0)
    out["dominant"] = f"{dom_type} ({dom_pct:.1f}% of {total_feat})"

    guide_by_hazard = _LAYER_METRIC_GUIDE.get(key, {})
    metric_guide = (
        guide_by_hazard.get(hazard_label)
        or next(iter(guide_by_hazard.values()), None)
        or _LAYER_EXTRA_CONTEXT.get(key, "")
    )
    city = site_info.get("city", "")
    country = site_info.get("country", "")
    location = ", ".join(p for p in [city, country] if p) or "the assessed location"

    old_prompt = _OLD_VECTOR_PROMPT.format(
        site_name=site_info.get("site_name", "the assessed site"),
        location=location,
        hazard=hazard_label,
        hazard_lower=hazard_label.lower(),
        layer_name=display_name,
        dom_type=dom_type,
        dom_pct=dom_pct,
        total_feat=total_feat,
        type_stats_json=json.dumps(type_stats, indent=2),
        metric_guide=metric_guide,
    )
    logger.info("    Calling Claude with OLD prompt...")
    out["before"] = _call_claude_with_old_prompt(old_prompt) or "(OLD Claude call failed)"

    logger.info("    Calling Claude with NEW prompt...")
    out["after"] = (
        _ask_claude_vector(display_name, hazard_label, type_stats, site_info, key)
        or "(NEW Claude call failed)"
    )
    return out


# ─────────────────────────────────────────────────────────────────────────────
# Word document builder
# ─────────────────────────────────────────────────────────────────────────────

def _build_word_doc(area: str, results: list[dict], output_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, RGBColor
    from docx.enum.table import WD_ALIGN_VERTICAL

    doc = Document()

    # Title page
    doc.add_heading("Appendix A — Impact Text BEFORE vs AFTER", level=0)
    doc.add_paragraph(f"Location: {area}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    doc.add_paragraph(
        "This document compares the LLM-generated impact text in Appendix A "
        "before and after the prompt rewrite for plain-English audience. For "
        "each data layer, both prompts use IDENTICAL raster statistics, "
        "susceptibility distributions and feature counts. The only thing "
        "that differs is the prompt instructions sent to Claude."
    )
    doc.add_paragraph(
        "Use this to validate that the new prompts produce text accessible "
        "to non-technical business stakeholders while preserving the same "
        "factual content."
    )
    doc.add_page_break()

    # Per-layer comparison pages
    for r in results:
        doc.add_heading(f"{r['layer_name']} ({r['layer_key']})", level=1)

        meta = doc.add_paragraph()
        meta.add_run("Hazard: ").bold = True
        meta.add_run(f"{r['hazard']}    ")
        if r.get("severity"):
            meta.add_run("Severity: ").bold = True
            meta.add_run(f"{r['severity']}    ")
        if r.get("dominant"):
            meta.add_run("Dominant: ").bold = True
            meta.add_run(f"{r['dominant']}")

        if r.get("skipped_reason"):
            warn = doc.add_paragraph()
            run = warn.add_run(f"⚠ Skipped: {r['skipped_reason']}")
            run.italic = True
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            doc.add_page_break()
            continue

        table = doc.add_table(rows=2, cols=2)
        table.style = "Light Grid"
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(3.25)

        # Header row
        hdr = table.rows[0].cells
        h0 = hdr[0].paragraphs[0].add_run("BEFORE — old technical prompt")
        h0.bold = True
        h0.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
        h1 = hdr[1].paragraphs[0].add_run("AFTER — new plain-English prompt")
        h1.bold = True
        h1.font.color.rgb = RGBColor(0x00, 0x66, 0x33)

        # Body row
        body = table.rows[1].cells
        body[0].text = r.get("before") or "(no output)"
        body[1].text = r.get("after")  or "(no output)"
        for cell in body:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        doc.add_page_break()

    doc.save(output_path)


# ─────────────────────────────────────────────────────────────────────────────
# CLI entry
# ─────────────────────────────────────────────────────────────────────────────

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Compare Appendix A impact text BEFORE vs AFTER the prompt rewrite.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--area",    default="Alagiyanallur", help="Site / area name")
    parser.add_argument("--city",    default="Virudhunagar",  help="City")
    parser.add_argument("--state",   default="Tamil Nadu",    help="State")
    parser.add_argument("--country", default="India",         help="Country")
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s  %(levelname)-8s  %(message)s",
        datefmt="%H:%M:%S",
    )
    # Silence the very chatty libraries
    logging.getLogger("anthropic").setLevel(logging.WARNING)
    logging.getLogger("httpx").setLevel(logging.WARNING)
    logging.getLogger("rasterio").setLevel(logging.WARNING)

    site_info = {
        "site_name": args.area,
        "city":      args.city,
        "state":     args.state,
        "country":   args.country,
    }
    cog_dir = config.COG_DIR

    # Compute overall site risk profiles from Input_Files GeoJSONs so the
    # AFTER prompt (which expects them) reflects the same calibration the
    # live pipeline would produce. Falls back gracefully if files missing.
    _inject_overall_risk_profiles(site_info, config.INPUT_FILES_DIR)

    logger.info("=" * 60)
    logger.info("  Appendix A tone comparison — Location: %s", args.area)
    logger.info("  COG dir: %s", cog_dir)
    logger.info("  Overall flood risk: %s", site_info.get("flood_overall"))
    logger.info("  Overall heat  risk: %s", site_info.get("heat_overall"))
    logger.info("=" * 60)

    results: list[dict] = []

    for (key, hazard, display_name, _fig_ref, _out_png, _cmap, _cbar, mode) in _LAYER_CATALOGUE:
        logger.info("Layer (raster): %s — %s", display_name, hazard)
        results.append(_process_raster_layer(key, hazard, display_name, mode, cog_dir, site_info))

    for (key, subdir_name, display_name, _fig_ref, _out_png) in _GEOJSON_CATALOGUE:
        logger.info("Layer (vector): %s", display_name)
        results.append(_process_vector_layer(key, subdir_name, display_name, cog_dir, site_info))

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    slug = args.area.replace(" ", "_").replace(",", "")
    output_path = config.REPORT_DATA_DIR / f"Comparison_{slug}_{timestamp}.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    _build_word_doc(args.area, results, output_path)

    ok_count   = sum(1 for r in results if r.get("before") and r.get("after") and not r.get("skipped_reason"))
    skip_count = sum(1 for r in results if r.get("skipped_reason"))
    logger.info("=" * 60)
    logger.info("  Done — %d layers compared, %d skipped", ok_count, skip_count)
    logger.info("  Output: %s", output_path)
    logger.info("=" * 60)


if __name__ == "__main__":
    main()
