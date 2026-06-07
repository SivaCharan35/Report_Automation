"""
Word Report Assembler.
Builds the final Asset Resilience Assessment .docx from the pipeline context.
Saved locally only — never uploaded to Azure (per design spec).

Public API:  build(context: dict) -> Path | None
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import config
from core.classifiers import (
    RISK_ORDER,
    classify_current_flood,
    classify_current_heat,
    classify_flood_ssp,
    classify_heat_ssp,
)
from core.geojson_utils import build_ssp_counts
from core.word_utils import (
    COL_BODY,
    COL_CAPTION,
    COL_HEADING,
    COL_SUBHEAD,
    COL_WHITE,
    HDR_BG_HEX,
    add_body,
    add_bullet,
    add_caption,
    add_header_footer,
    add_heading,
    add_image,
    add_page_break,
    add_run,
    add_subheading,
    insert_2x2_chart_grid,
    set_default_font,
    set_doc_margins,
    style_cell,
)

logger = logging.getLogger(__name__)

# ── SSP horizons (re-used in this file; sourced from config) ──────────────────
_SSP_HORIZONS = config.SSP_HORIZONS

# ── Risk-level row background colours ────────────────────────────────────────
_ROW_BG: dict[str, str] = {
    "Very Low":  "E2EFDA",
    "Low":       "DDEBF7",
    "Moderate":  "FFF2CC",
    "High":      "FCE4D6",
    "Very High": "F4CCCC",
}

# ── Hazard accent colours ─────────────────────────────────────────────────────
_FLOOD_ACCENT = RGBColor(0x21, 0x96, 0xF3)
_HEAT_ACCENT  = RGBColor(0xFF, 0x57, 0x22)


# ═════════════════════════════════════════════════════════════════════════════
# PAGE BUILDERS (each builds one logical page / section)
# ═════════════════════════════════════════════════════════════════════════════

def _cover_page(doc: Document, ctx: dict) -> None:
    """Cover page + Table of Contents."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p,
        "The document is confidential and proprietary to Resilience AI  |  1 | P a g e",
        size=8, italic=True, color=RGBColor(0x60, 0x60, 0x60))

    for _ in range(3):
        doc.add_paragraph()

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_title, "Asset Resilience Assessment", bold=True, size=26, color=COL_HEADING)
    doc.add_paragraph()

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_sub, "Climate Risk Assessment Report", size=14, color=COL_SUBHEAD)

    for _ in range(2):
        doc.add_paragraph()

    for label, value, dyn in [
        ("Prepared by:", "Resilience AI Solutions Private. Ltd.", False),
        ("Area covered:", ctx["area_covered_full"], True),
        ("Prepared for:", ctx["client"], True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f"{label} ", bold=True, size=12, color=COL_HEADING)
        add_run(p, value, size=12, dynamic=dyn)

    for _ in range(2):
        doc.add_paragraph()

    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p_line, "─" * 70, size=10, color=RGBColor(0xCC, 0xCC, 0xCC))
    doc.add_paragraph()

    # Table of Contents
    add_run(doc.add_paragraph(), "Table of Contents", bold=True, size=14, color=COL_HEADING)
    doc.add_paragraph()

    toc: list[tuple[str, str]] = [
        ("",     "Executive Summary"),
        ("1.",   "Introduction"),
        ("2.",   "Methodology"),
        ("3.",   "Hazard Risk Assessment"),
        ("3.1",  "  Hazard Risk Assessment Overview"),
        ("3.2",  "  SSP Analysis"),
        ("4.",   "Impact Scale"),
        ("5.",   "Historical Trends"),
    ]
    rf = ctx["risk_for"]
    if rf in ("Flood", "Both"):
        toc.append(("5.1", "  Flood"))
    if rf in ("Heat", "Both"):
        lbl = "5.2" if rf == "Both" else "5.1"
        toc.append((lbl, "  Heat"))

    toc.append(("6.", "Influencing Factors"))
    if rf in ("Flood", "Both"):
        toc.append(("6.1", "  Flood"))
    if rf in ("Heat", "Both"):
        lbl = "6.2" if rf == "Both" else "6.1"
        toc.append((lbl, "  Heat"))

    toc.append(("7.", "Appendices"))
    toc.append(("A.", "  Appendix A: Risk Assessment Data"))
    toc.append(("B.", "  Appendix B: SSP Scenarios"))

    for num, title in toc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        add_run(p, f"{num}  {title}" if num else f"       {title}",
                bold=(not num), size=10, color=COL_BODY)


def _intro_methodology(doc: Document, ctx: dict) -> None:
    """Executive Summary, Section 1 — Introduction, Section 2 — Methodology."""
    add_page_break(doc)

    # Executive Summary
    add_heading(doc, "Executive Summary", size=13)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, "This report presents the detailed climate risk assessment for ", size=10)
    add_run(p, ctx["site_name"],     size=10, dynamic=True)
    add_run(p, ", ",                 size=10)
    add_run(p, ctx["site_location"], size=10, dynamic=True)
    hazard_phrase = {
        "Both":  " focusing on Flood and Heat Assessment.",
        "Flood": " focusing on Flood Assessment.",
        "Heat":  " focusing on Heat Assessment.",
    }[ctx["risk_for"]]
    add_run(p, hazard_phrase + (
        " We have used Resilience360™, a proprietary enterprise software suite to "
        "perform this assessment. The assessment analyses severity of risk for each "
        "building and open area in the identified area of interest (referred to as "
        "'location' from here on)."
    ), size=10)

    # Section 1
    add_heading(doc, "1.  Introduction", size=13)
    add_body(doc,
        "ResSolv™ is an advanced climate risk and resilience quantification platform "
        "developed by Resilience AI Solutions. It integrates multi-source geospatial "
        "data, physics-based modelling, and machine learning to deliver building-level "
        "climate risk scores across a wide range of hazards. The platform enables "
        "organisations to understand, quantify, and manage the physical climate risks "
        "associated with their assets under both current conditions and future climate "
        "scenarios."
    )
    add_body(doc, "The key objectives of this assessment are:")
    add_bullet(doc,
        "To quantify the current and projected flood and heat risk exposure at the "
        "building and asset level for the identified area of interest.")
    add_bullet(doc,
        "To evaluate risk trajectories across Near-Term (2025–2040), Medium-Term "
        "(2041–2060), and Long-Term (2081–2100) horizons under multiple RCP and SSP scenarios.")
    add_bullet(doc,
        "To support informed decision-making on climate adaptation, asset protection, "
        "and operational resilience planning.")

    # Section 2
    add_heading(doc, "2.  Methodology", size=13)
    add_body(doc, ctx.get("methodology_text", "").split("SECTION 2")[0] or (
        "The assessment is conducted using the ResSolv™ platform, which evaluates "
        "climate risk through an integrated multi-hazard framework. Risk scores are "
        "generated at the individual building level using a composite scoring methodology "
        "that accounts for 20 parameters spanning physical exposure, hazard intensity, "
        "sensitivity, and adaptive capacity."
    ))


def _risk_overview(doc: Document, ctx: dict) -> None:
    """Section 3.1 — Hazard Risk Assessment Overview (legend + risk maps)."""
    add_page_break(doc)
    add_heading(doc, "3.  Hazard Risk Assessment", size=13)
    add_subheading(doc, "3.1  Hazard Risk Assessment Overview", size=11)

    # Site bullets
    bullets = [
        ("Location: ",                      ctx["area_covered_full"]),
        ("Total buildings assessed: ",      f"{ctx.get('total_buildings', 0):,}"),
        ("Total area coverage: approximately ", f"{ctx.get('aoi_area', 0)} sq. miles"),
        ("Climate events assessed: ",       ctx["risk_for"]),
    ]
    for label, value in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_run(p, label, size=10)
        add_run(p, value, size=10, dynamic=True)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p,
        "The following section presents the hazard risk assessment results for ", size=10)
    add_run(p, ctx["site_name"], size=10, dynamic=True)
    add_run(p,
        ". Each building and open area within the defined area of interest has been "
        "individually assessed for flood and heat risk using the ResSolv™ platform. "
        "Risk scores are computed on a scale of 1 to 5, where higher scores indicate "
        "greater exposure and potential impact.", size=10)

    # Risk legend table
    add_subheading(doc, "Risk Score Legend", size=10, color=COL_HEADING)
    _risk_legend_table(doc)
    doc.add_paragraph()

    # Flood risk map
    if ctx["risk_for"] in ("Flood", "Both") and ctx.get("flood_risk_map_path"):
        add_subheading(doc, "Flood", size=11, color=_FLOOD_ACCENT)
        add_image(doc, ctx["flood_risk_map_path"], width=Inches(5.5),
                  caption=f"Figure 3.1 — Flood Risk Map: {ctx['site_name']}")

    # Heat risk map (new page if both)
    if ctx["risk_for"] in ("Heat", "Both") and ctx.get("heat_risk_map_path"):
        if ctx["risk_for"] == "Both":
            add_page_break(doc)
        add_subheading(doc, "Heat", size=11, color=_HEAT_ACCENT)
        add_image(doc, ctx["heat_risk_map_path"], width=Inches(5.5),
                  caption=f"Figure 3.2 — Heat Risk Map: {ctx['site_name']}")


def _risk_legend_table(doc: Document) -> None:
    LEGEND = [
        ("1", "VERY LOW",  "4CAF50"),
        ("2", "LOW",       "8BC34A"),
        ("3", "MODERATE",  "FFC107"),
        ("4", "HIGH",      "FF5722"),
        ("5", "VERY HIGH", "B71C1C"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    style_cell(tbl.rows[0].cells[0], "Score",      bold=True, size=9, bg_color=HDR_BG_HEX, font_color=COL_WHITE)
    style_cell(tbl.rows[0].cells[1], "Risk Level", bold=True, size=9, bg_color=HDR_BG_HEX, font_color=COL_WHITE)
    for score, label, hex_c in LEGEND:
        row = tbl.add_row().cells
        style_cell(row[0], score, bold=True, size=9, bg_color=hex_c, font_color=COL_WHITE)
        style_cell(row[1], label, bold=True, size=9, bg_color=hex_c, font_color=COL_WHITE)


def _risk_summary_table(doc: Document, ctx: dict) -> None:
    """Risk Level Summary (flood counts + heat counts side by side)."""
    doc.add_paragraph()
    add_subheading(doc, "Risk Level Summary", size=11, color=COL_HEADING)

    flood_c = ctx.get("flood_risk_counts", {})
    heat_c  = ctx.get("heat_risk_counts",  {})
    rl_info = [
        ("Very Low",  "1", "4CAF50"),
        ("Low",       "2", "8BC34A"),
        ("Moderate",  "3", "FFC107"),
        ("High",      "4", "FF5722"),
        ("Very High", "5", "B71C1C"),
    ]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hdr in enumerate(["RISK LEVEL", "SCORE", "FLOOD", "HEAT"]):
        style_cell(tbl.rows[0].cells[i], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)
    for label, score, hex_c in rl_info:
        row = tbl.add_row().cells
        style_cell(row[0], label, bold=True, size=9, bg_color=hex_c, font_color=COL_WHITE)
        style_cell(row[1], score, bold=True, size=9, bg_color=hex_c, font_color=COL_WHITE)
        style_cell(row[2], str(flood_c.get(label, "")), size=9)
        style_cell(row[3], str(heat_c.get(label,  "")), size=9)
    doc.add_paragraph()


def _ssp_word_table(
    doc: Document,
    geojson_data: dict,
    classify_today_fn,
    classify_ssp_fn,
    heading: str,
    color: RGBColor,
) -> None:
    """Native Word SSP projection table (11 cols × 7 rows)."""
    add_subheading(doc, heading, size=10, color=color)

    today_counts, cols, per_col = build_ssp_counts(
        geojson_data, classify_today_fn, classify_ssp_fn, _SSP_HORIZONS
    )

    tbl = doc.add_table(rows=7, cols=11)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths
    col_widths = [Inches(1.1), Inches(0.60)] + [Inches(0.47)] * 9
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_widths[ci]

    # Row 0 — merged time-horizon headers
    tbl.cell(0, 2).merge(tbl.cell(0, 4))
    tbl.cell(0, 5).merge(tbl.cell(0, 7))
    tbl.cell(0, 8).merge(tbl.cell(0, 10))
    for ci, txt in [(0, ""), (1, ""),
                    (2, "Near Term [2025-2040]"),
                    (5, "Medium Term [2041-2060]"),
                    (8, "Long Term [2081-2100]")]:
        style_cell(tbl.cell(0, ci), txt, bg_color=HDR_BG_HEX,
                   font_color=COL_WHITE, bold=(ci > 1), size=9)

    # Row 1 — column sub-headers
    for ci, lbl in enumerate(["Risk Level", "2025\n(Today)",
                               "SSP 2.6", "SSP 4.5", "SSP 8.5",
                               "SSP 2.6", "SSP 4.5", "SSP 8.5",
                               "SSP 2.6", "SSP 4.5", "SSP 8.5"]):
        style_cell(tbl.rows[1].cells[ci], lbl,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE, bold=True, size=8)

    # Rows 2-6 — data
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A
    for ri, rl in enumerate(RISK_ORDER):
        row = tbl.rows[2 + ri]
        bg  = _ROW_BG[rl]
        style_cell(row.cells[0], rl, bg_color=bg, bold=True, size=9,
                   align=_A.LEFT)
        tv = today_counts.get(rl, 0)
        style_cell(row.cells[1], str(tv) if tv else "", bg_color=bg, size=9)
        for ci, col_key in enumerate(cols):
            cnt = per_col.get(col_key, {}).get(rl, 0)
            style_cell(row.cells[2 + ci], str(cnt) if cnt else "", bg_color=bg, size=9)

    doc.add_paragraph()


def _ssp_section(doc: Document, ctx: dict) -> None:
    """Section 3.2 — SSP Analysis."""
    add_page_break(doc)
    add_heading(doc, "3.2  SSP Analysis", size=13)
    add_body(doc,
        "The SSP (Shared Socioeconomic Pathway) analysis evaluates how projected changes "
        "in greenhouse gas emissions influence the future intensity of flood and heat risk "
        "at the site. Risk projections are shown for Near Term (2025–2040), Medium Term "
        "(2041–2060), and Long Term (2081–2100) under SSP 2.6, 4.5, and 8.5 scenarios."
    )

    if ctx["risk_for"] in ("Flood", "Both") and ctx.get("flood_data"):
        _ssp_word_table(doc, ctx["flood_data"],
                        classify_current_flood, classify_flood_ssp,
                        "Flood — SSP Projections", _FLOOD_ACCENT)
        add_caption(doc, "Table 3.1 — Flood Risk by SSP Scenario and Time Horizon")
        doc.add_paragraph()

    if ctx["risk_for"] in ("Heat", "Both") and ctx.get("heat_data"):
        _ssp_word_table(doc, ctx["heat_data"],
                        classify_current_heat, classify_heat_ssp,
                        "Heat — SSP Projections", _HEAT_ACCENT)
        add_caption(doc, "Table 3.2 — Heat Risk by SSP Scenario and Time Horizon")


def _impact_section(doc: Document, ctx: dict) -> None:
    """Section 4 — Impact Scale."""
    add_page_break(doc)
    add_heading(doc, "4.  Impact Scale", size=13)
    add_subheading(doc, "Aggregated Risk Scale for Structural and Community Impacts",
                   size=11, color=COL_HEADING)
    add_body(doc,
        "The following table presents the aggregated risk scale used to assess the "
        "structural, community, and operational impacts associated with each risk level "
        "for flood and heatwave hazards at the site."
    )

    impact_data = ctx.get("impact_data", [])
    tbl = doc.add_table(rows=len(impact_data) + 1, cols=4)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.8)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    for i, hdr in enumerate(["Risk Level", "Flood Impact", "Heatwave Impact", "Operational Impact"]):
        style_cell(tbl.rows[0].cells[i], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A
    for ri, item in enumerate(impact_data):
        row = tbl.rows[ri + 1]
        hex_c = item["color_hex"].lstrip("#")
        style_cell(row.cells[0], item["risk_level"], bold=True, size=8,
                   bg_color=hex_c, font_color=COL_WHITE)
        style_cell(row.cells[1], item["flood_impact"],        size=8, align=_A.LEFT)
        style_cell(row.cells[2], item["heat_impact"],         size=8, align=_A.LEFT)
        style_cell(row.cells[3], item["operational_impact"],  size=8, align=_A.LEFT)
    doc.add_paragraph()


def _historical_section(doc: Document, ctx: dict) -> None:
    """Section 5 — Historical Trends."""
    add_page_break(doc)
    add_heading(doc, "5.  Historical Trends", size=13)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_run(p, "This section presents historical climate trend analysis for ", size=10)
    add_run(p, ctx["site_name"],     size=10, dynamic=True)
    add_run(p, ", ",                 size=10)
    add_run(p, ctx["site_location"], size=10, dynamic=True)
    add_run(p,
        ". These historical trends provide the empirical basis for projecting future "
        "climate risk trajectories and validating the risk scores generated by the "
        "ResSolv™ platform.", size=10)

    hist       = ctx.get("historical_charts", {})
    flood_charts = hist.get("flood", [])
    heat_charts  = hist.get("heat",  [])
    risk_for     = ctx["risk_for"]

    if risk_for in ("Flood", "Both") and flood_charts:
        add_subheading(doc, "5.1  Flood", size=11, color=_FLOOD_ACCENT)
        add_body(doc,
            "The following charts illustrate historical flood-relevant parameters including "
            "precipitation frequency, intensity, runoff, and total weekly precipitation."
        )
        insert_2x2_chart_grid(doc, flood_charts)
        doc.add_paragraph()

    if risk_for in ("Heat", "Both") and heat_charts:
        section_label = "5.2" if risk_for == "Both" else "5.1"
        add_subheading(doc, f"{section_label}  Heat", size=11, color=_HEAT_ACCENT)
        add_body(doc,
            "The following charts illustrate historical heat-relevant parameters including "
            "heatwave frequency, maximum air temperature, land surface temperature (LST), "
            "and the heat index."
        )
        insert_2x2_chart_grid(doc, heat_charts)


# ── Influencing Factors — static data ────────────────────────────────────────

# Risk level colours matching the PNG map legend
_IF_RISK_HEX: dict[str, str] = {
    "Very Low": "4CEB34",   # bright green
    "Low":      "EBEB34",   # bright yellow
    "High":     "EB8F34",   # orange
    "Very High":"EB3434",   # red
}

# Risk summary table: per-hazard static interpretation text per risk level
_FLOOD_IF_INTERP: dict[str, str] = {
    "Very Low":  "",
    "Low":       "Low TWI values with sparse vegetation (low NDVI), suggesting limited "
                 "water retention capacity but reduced runoff buffering.",
    "Moderate":  "Localized zones of elevated TWI, partially offset by high NDVI.",
    "High":      "High TWI values combined with moderate distance from the river channel, "
                 "leading to significant water accumulation potential.",
    "Very High": "Immediate proximity to the river, where fluvial flooding dominates "
                 "despite vegetative cover (high NDVI) and relatively low TWI.",
}

_HEAT_IF_INTERP: dict[str, str] = {
    "Very Low":  "Low LST values with sparse vegetation (low NDVI), suggesting limited "
                 "latent heat build-up.",
    "Low":       "Localized zones of elevated LST, partially offset by high NDVI.",
    "Moderate":  "High LST values combined with increased presence of built area (high NDBI).",
    "High":      "High NDBI values (low density built area and barren land) with reduced "
                 "green cover become hotspots for LST hotspot formation.",
    "Very High": "Very high levels of NDBI (indicating high density-built area) and very "
                 "low levels of NDVI (barren land) indicate heightened levels of heat "
                 "stress and impacts from increased LST.",
}

# Color-class legend tables — (risk_level_hex, col_a_value, col_b_value, interpretation)
_NDVI_TWI_CLASSES = [
    ("EBEB34", "Low",  "Low",  "Sparse vegetation with low water accumulation; areas with "
                               "limited green cover and good drainage, indicating minimal flood risk."),
    ("EB3434", "High", "Low",  "Sparse vegetation with high water accumulation; bare or less "
                               "vegetated land in low-lying areas where water tends to collect, "
                               "indicating high flood vulnerability."),
    ("4CEB34", "Low",  "High", "Dense vegetation with low water accumulation; well-vegetated "
                               "areas on slopes or elevated terrain with good drainage and "
                               "minimal waterlogging."),
    ("EB8F34", "High", "High", "Dense vegetation with high water accumulation; vegetated areas "
                               "in depressions where water naturally accumulates, indicating "
                               "moderate to high flood susceptibility."),
]

_NDBI_DEM_CLASSES = [
    ("EBEB34", "Low",  "Low",  "Semi-open, low-lying areas with moderate flood susceptibility."),
    ("EB3434", "High", "Low",  "Dense built-up in low elevation; critical flood-prone pockets "
                               "due to poor drainage."),
    ("4CEB34", "Low",  "High", "Elevated but less urbanised; relatively safe zones with "
                               "natural drainage."),
    ("EB8F34", "High", "High", "High built-up on elevated land; less flood risk locally but "
                               "contributes to downstream runoff."),
]

_LST_NDVI_CLASSES = [
    ("EBEB34", "Low",  "Low",  "Areas with low density vegetation, indicating moist ground "
                               "resulting in low LST values."),
    ("4CEB34", "Low",  "High", "Areas with high density vegetation, indicating higher levels "
                               "of moisture and low levels of LST."),
    ("EB3434", "High", "Low",  "Likely barren areas with low levels of vegetation and high "
                               "levels of surface temperature."),
    ("EB8F34", "High", "High", "Built areas with high surface temperatures along with vegetation."),
]


# ── Influencing Factors — helper table builders ───────────────────────────────

def _if_risk_table(
    doc: Document,
    risk_counts: dict,
    interp_map: dict,
    col_header: str,
) -> None:
    """
    Render a 3-column risk summary table:
      RISK LEVEL  |  SCORE  |  <col_header> INTERPRETATION
    Score is pulled from risk_counts; interpretation from interp_map (static).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A

    _RL_INFO = [
        ("Very Low",  "1", "4CAF50"),
        ("Low",       "2", "8BC34A"),
        ("Moderate",  "3", "FFC107"),
        ("High",      "4", "FF5722"),
        ("Very High", "5", "B71C1C"),
    ]
    col_widths = [Inches(1.2), Inches(0.6), Inches(4.7)]

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, cell in enumerate(tbl.rows[0].cells):
        cell.width = col_widths[ci]

    for ci, hdr in enumerate(["RISK LEVEL", "SCORE", f"{col_header.upper()} INTERPRETATION"]):
        style_cell(tbl.rows[0].cells[ci], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    for label, score, hex_c in _RL_INFO:
        row = tbl.add_row().cells
        for ci, cell in enumerate(row):
            cell.width = col_widths[ci]
        style_cell(row[0], label, bold=True, size=9, bg_color=hex_c, font_color=COL_WHITE)
        cnt = risk_counts.get(label, 0)
        style_cell(row[1], str(cnt) if cnt else "—", size=9)
        style_cell(row[2], interp_map.get(label, ""), size=9, align=_A.LEFT)

    doc.add_paragraph()


def _if_color_class_table(
    doc: Document,
    col_a_name: str,
    col_b_name: str,
    rows: list,
) -> None:
    """
    Render the 4-row colour-class legend table:
      [Colour swatch]  |  col_a_name  |  col_b_name  |  Interpretation
    rows = list of (hex_color, a_val, b_val, interpretation).
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A

    col_widths = [Inches(0.35), Inches(0.75), Inches(0.75), Inches(4.65)]

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, cell in enumerate(tbl.rows[0].cells):
        cell.width = col_widths[ci]

    for ci, hdr in enumerate(["Colour\nClass", col_a_name, col_b_name, "Interpretation"]):
        style_cell(tbl.rows[0].cells[ci], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    for hex_c, a_val, b_val, interp in rows:
        row = tbl.add_row().cells
        for ci, cell in enumerate(row):
            cell.width = col_widths[ci]
        style_cell(row[0], "", bg_color=hex_c)             # coloured swatch, no text
        style_cell(row[1], a_val, bold=True, size=9)
        style_cell(row[2], b_val, bold=True, size=9)
        style_cell(row[3], interp, size=9, align=_A.LEFT)

    doc.add_paragraph()


# ── Section 6 assembler ───────────────────────────────────────────────────────

def _influencing_factors_section(doc: Document, ctx: dict) -> None:
    """Section 6 — Influencing Factors."""
    if_paths   = ctx.get("influencing_factor_paths", {})
    flood_maps = if_paths.get("flood", [])
    heat_maps  = if_paths.get("heat",  [])

    if not flood_maps and not heat_maps:
        return

    flood_counts = ctx.get("flood_risk_counts", {})
    heat_counts  = ctx.get("heat_risk_counts",  {})
    risk_for     = ctx["risk_for"]

    add_page_break(doc)
    add_heading(doc, "6.  Influencing Factors", size=13)

    # ── 6.1  Flood ────────────────────────────────────────────────────────────
    if risk_for in ("Flood", "Both") and flood_maps:
        add_subheading(doc, "6.1  Flood", size=11, color=_FLOOD_ACCENT)
        add_body(doc,
            "Flood risk arises from the combined effects of natural conditions and human "
            "activities, leading to variations in exposure and vulnerability across different "
            "areas. Understanding these patterns is essential for identifying zones that are "
            "more susceptible to flooding and for planning targeted mitigation strategies. "
            "Spatial variations can give insights into areas of concern and support effective "
            "disaster resilience planning."
        )

        # Risk summary table
        _if_risk_table(doc, flood_counts, _FLOOD_IF_INTERP, "Flood")

        # Map 1 — NDVI × TWI
        ndvi_twi = next((p for p in flood_maps if "ndvi_twi" in p.name), None)
        if ndvi_twi:
            add_subheading(doc, "Influence of NDVI (Vegetation) versus TWI (Wetness)",
                           size=10, color=COL_HEADING)
            add_image(doc, ndvi_twi, width=Inches(5.5),
                      caption=f"Figure 6.1  Map showing influence of NDVI vs TWI")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_run(p, "\u25ba ", bold=True, size=10, color=COL_HEADING)
            add_run(p,
                "Areas with high TWI values show greater water accumulation potential. "
                "Zones combining low NDVI (reduced vegetation) with high TWI face the highest "
                "flood exposure, as limited vegetation does little to mitigate water retention "
                "and surface runoff.",
                size=10)
            _if_color_class_table(doc, "NDVI\n(Vegetation)", "TWI\n(Wetness)",
                                  _NDVI_TWI_CLASSES)

        # Map 2 — NDBI × DEM
        ndbi_dem = next((p for p in flood_maps if "ndbi_dem" in p.name), None)
        if ndbi_dem:
            add_subheading(doc, "Influence of NDBI (Built-up) versus DEM (Elevation)",
                           size=10, color=COL_HEADING)
            add_image(doc, ndbi_dem, width=Inches(5.5),
                      caption=f"Figure 6.2  Map showing influence of NDBI vs Elevation")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_run(p, "\u25ba ", bold=True, size=10, color=COL_HEADING)
            add_run(p,
                "Low-elevation areas with high built-up density show elevated flood vulnerability "
                "due to impervious surfaces and reduced drainage capacity. Locations at lower "
                "elevations are at increased risk of waterlogging if drainage is overwhelmed.",
                size=10)
            _if_color_class_table(doc, "Built-up\n(NDBI)", "Elevation\n(DEM)",
                                  _NDBI_DEM_CLASSES)

    # ── 6.2  Heat ────────────────────────────────────────────────────────────
    if risk_for in ("Heat", "Both") and heat_maps:
        section_lbl = "6.2" if risk_for == "Both" else "6.1"
        add_subheading(doc, f"{section_lbl}  Heat", size=11, color=_HEAT_ACCENT)
        add_body(doc,
            "Heatwave risk emerges from climatic conditions resulting in varying levels of "
            "exposure and vulnerability across different regions. Understanding these spatial "
            "patterns is crucial for identifying areas that are more susceptible to extreme "
            "heat events and for developing targeted adaptation measures. Analysing spatial "
            "variations helps pinpoint vulnerable communities and supports effective heat "
            "action planning and public health interventions."
        )

        # Map — LST × NDVI
        lst_ndvi = next((p for p in heat_maps if "lst_ndvi" in p.name), None)
        if lst_ndvi:
            add_subheading(doc, "Influence of LST versus NDVI",
                           size=10, color=COL_HEADING)
            add_image(doc, lst_ndvi, width=Inches(5.5),
                      caption=f"Figure 6.3  LST vs NDVI map")
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_run(p, "\u25ba ", bold=True, size=10, color=COL_HEADING)
            add_run(p,
                "Areas with high land surface temperature and low vegetation coverage represent "
                "urban heat islands and face the highest heat stress. Dense built-up zones with "
                "minimal tree cover experience the most intense surface heating, while "
                "well-vegetated areas show significantly reduced thermal exposure.",
                size=10)

        # Risk summary table (heat)
        _if_risk_table(doc, heat_counts, _HEAT_IF_INTERP, "Heat")

        # Colour-class table
        if lst_ndvi:
            _if_color_class_table(doc, "LST\n(Temperature)", "NDVI\n(Vegetation)",
                                  _LST_NDVI_CLASSES)


# ── Appendices — static text (mirrors scripts/7_appendices.py) ────────────────

_LAYER_APPENDIX_TEXT: dict[str, dict[str, str]] = {
    "dem": {
        "description": (
            "The Digital Elevation Model (DEM) represents the topographic surface of the study area, "
            "derived from high-resolution satellite data. Elevation data captures natural terrain "
            "features that govern surface water flow patterns and flood accumulation zones. The DEM "
            "also serves as a foundational input for deriving secondary terrain attributes such as "
            "slope, flow direction, and the Topographic Wetness Index."
        ),
        "impact": (
            "Low-lying areas are naturally more susceptible to flood inundation as water flows and "
            "accumulates in topographic depressions. High-elevation terrain serves as a natural flood "
            "buffer, while valleys and flat low-elevation zones face heightened risk from both riverine "
            "and pluvial flooding. Structures in low-elevation zones with limited drainage infrastructure "
            "face compounded risk during prolonged or intense rainfall events."
        ),
    },
    "twi": {
        "description": (
            "The Topographic Wetness Index (TWI) is derived from the DEM and quantifies the tendency "
            "of each landscape position to accumulate water, based on the upstream contributing area "
            "and local slope gradient. Higher TWI values indicate zones where water naturally converges "
            "and saturates the soil, irrespective of precipitation intensity."
        ),
        "impact": (
            "Zones with high TWI values represent natural water convergence areas where surface runoff "
            "aggregates. These areas face elevated flood susceptibility during intense rainfall events, "
            "as terrain geometry drives water towards these focal points regardless of surface land "
            "cover. Combined with proximity to water bodies, high-TWI regions are consistently among "
            "the first affected during flood events."
        ),
    },
    "impervious": {
        "description": (
            "The Impervious Surface Cover layer is a binary classification derived from high-resolution "
            "satellite imagery, mapping every pixel as either fully impervious (value = 1: concrete, "
            "asphalt, rooftops) or pervious (value = 0: soil, vegetation, open water). Impervious "
            "surfaces generate near-total surface runoff during rainfall events, as water cannot "
            "infiltrate and must drain entirely via surface pathways. The proportion of impervious "
            "cover within the study area is a primary determinant of peak runoff volume and velocity."
        ),
    },
    "ndvi": {
        "description": (
            "The Normalised Difference Vegetation Index (NDVI) measures vegetation density and health "
            "using near-infrared (NIR) and red (RED) spectral bands from satellite imagery. NDVI values "
            "range from -1 to +1, with higher positive values (> 0.5) indicating dense, healthy "
            "vegetation such as forests and croplands, while values near zero or negative indicate "
            "bare soil, impervious surfaces, or water bodies."
        ),
        "impact": (
            "Vegetation plays a critical role in moderating heat risk through evapotranspiration and "
            "shading. Dense vegetation (high NDVI) reduces land surface temperatures and creates cooler "
            "microclimates. Areas with sparse or absent vegetation (low NDVI) are more exposed to solar "
            "radiation, contributing to elevated surface temperatures and heat stress, particularly "
            "during prolonged heatwave events."
        ),
    },
    "ndbi": {
        "description": (
            "The Normalised Difference Built-Up Index (NDBI) identifies impervious surfaces and urban "
            "built-up areas using shortwave infrared (SWIR) and near-infrared (NIR) spectral bands "
            "from satellite imagery. Positive NDBI values indicate the presence of built-up or "
            "impervious surfaces, while negative values correspond to vegetation or water bodies."
        ),
        "impact": (
            "High NDBI areas represent dense urban development where impervious surfaces absorb and "
            "retain solar radiation, directly contributing to urban heat island effects. The replacement "
            "of natural surfaces with concrete and asphalt significantly elevates land surface "
            "temperatures and amplifies heat stress during heatwave periods. Urban densification trends "
            "make NDBI a forward-looking indicator of increasing heat vulnerability."
        ),
    },
    "lst": {
        "description": (
            "Land Surface Temperature (LST) is derived from thermal infrared satellite imagery and "
            "represents the radiative skin temperature of the Earth's surface as observed from space. "
            "LST integrates the combined thermal effects of land cover, solar radiation, vegetation "
            "density, soil moisture, and urban morphology, making it a comprehensive indicator of "
            "surface heat exposure."
        ),
        "impact": (
            "Elevated LST values directly indicate areas experiencing thermal stress, particularly "
            "during heatwave events. Urban areas and bare ground consistently exhibit higher LST due "
            "to thermal mass accumulation and reduced evaporative cooling. Prolonged exposure to high "
            "LST increases heat-related health risks, infrastructure thermal loading, and peak energy "
            "demand. Vegetated zones and water bodies maintain lower temperatures through natural "
            "cooling mechanisms."
        ),
    },
    "lulc": {
        "description": (
            "Land Use / Land Cover (LULC) classification maps the Earth's surface into distinct "
            "functional categories including built-up areas, vegetation types, water bodies, and "
            "agricultural land. The classification is derived from multi-spectral satellite imagery "
            "using supervised machine learning algorithms and validated against ground-truth data."
        ),
        "impact": (
            "LULC fundamentally determines the thermal and hydrological behaviour of the landscape. "
            "Built-up areas and bare ground increase heat risk through surface albedo reduction and "
            "thermal mass effects, while also increasing flood risk by reducing infiltration capacity "
            "and elevating surface runoff coefficients. Natural vegetation and water bodies provide "
            "critical cooling benefits and flood mitigation through evapotranspiration and natural "
            "infiltration."
        ),
    },
    "roads": {
        "description": (
            "The road network layer maps all classified road and path features within the study area, "
            "sourced from OpenStreetMap (OSM) vector data. Road infrastructure is a critical component "
            "of climate risk assessment as it governs evacuation routes, emergency access corridors, "
            "and economic connectivity before, during, and after climate events. Roads also contribute "
            "to impervious surface cover, which directly affects surface runoff volumes and urban "
            "heat retention."
        ),
    },
    "waterline": {
        "description": (
            "The waterways and water bodies layer maps all surface water features within the study "
            "area, including rivers, canals, reservoirs, ponds, and detention basins, sourced from "
            "the National Hydrography Dataset (NHD). Surface water networks are primary determinants "
            "of flood pathway dynamics: active stream channels convey floodwaters, while engineered "
            "canals and ditches provide controlled drainage. Proximity to these features is a direct "
            "indicator of fluvial and pluvial flood exposure."
        ),
    },
}

# ── Appendix A conclusion template (hazard_types filled at render time) ──────

_APPENDIX_CONCLUSION_TEMPLATE = (
    "Each of these layers contributes to the accuracy and reliability of the climate risk "
    "assessment. By integrating multi-variate parameters, Resilience360\u2122 is able to produce "
    "detailed and localised risk assessments for {hazard_types}. These insights guide "
    "intervention strategies to mitigate the identified risks in the specified area of interest."
)

_HAZARD_LABEL: dict[str, str] = {
    "Flood": "flood risk",
    "Heat":  "heat risk",
    "Both":  "flood and heat risk",
}

# ── Appendix B — structured SSP content for Word rendering ───────────────────
# Each entry is (type, text):
#   "body"        → add_body paragraph
#   "ssp_heading" → bold subheading for each SSP label
#   "bullet"      → List Bullet paragraph

_SSP_APPENDIX_CONTENT: list[tuple[str, str]] = [
    ("body",
     "The Shared Socioeconomic Pathways (SSPs) are global scenarios developed by the "
     "Intergovernmental Panel on Climate Change (IPCC) for its Sixth Assessment Report (AR6). "
     "They describe possible future trajectories of societal development, economic growth, "
     "population dynamics, technology, and energy use. Unlike the Representative Concentration "
     "Pathways (RCPs), which focus on greenhouse gas concentrations, SSPs emphasise the "
     "socioeconomic factors that influence emissions and the capacity to mitigate or adapt to "
     "climate change."),
    ("body",
     "Each SSP is identified by a number that represents the radiative forcing (the change in "
     "energy balance in the Earth's atmosphere due to greenhouse gases) expected by the year "
     "2100, measured in watts per square metre (W/m\u00b2). There are three established SSP "
     "scenarios:"),
    ("ssp_heading", "SSP 2.6"),
    ("bullet",
     "Description: This is a low-emission scenario that assumes significant mitigation efforts. "
     "Gradual move towards sustainability and environmental respect; increasing action towards "
     "Sustainable Development Goals."),
    ("bullet", "Radiative Forcing: Stabilises around 2.6 W/m\u00b2 by 2100."),
    ("bullet",
     "Climate Implications: Global temperatures are expected to rise by approximately "
     "1.3\u20132.4\u00b0C by 2100."),
    ("ssp_heading", "SSP 4.5"),
    ("bullet",
     "Description: This scenario assumes moderate mitigation efforts. Similar to the past; "
     "unevenly distributed; slow progress towards SDGs."),
    ("bullet", "Radiative Forcing: Stabilises around 4.5 W/m\u00b2 by 2100."),
    ("bullet",
     "Climate Implications: Global temperatures are expected to rise by approximately "
     "2.1\u20133.5\u00b0C by 2100."),
    ("ssp_heading", "SSP 8.5"),
    ("bullet",
     "Description: This is a high-emission scenario often referred to as \u2018business-as-usual\u2019, "
     "where no significant changes are made to curb emissions. Resource-intensive lifestyles and "
     "industries; high investment in health and education; dependence on technological solutions "
     "at the expense of the environment."),
    ("bullet", "Radiative Forcing: Reaches 8.5 W/m\u00b2 by 2100."),
    ("bullet",
     "Climate Implications: This scenario could lead to a global temperature rise of more than "
     "3.3\u20135.7\u00b0C by 2100, resulting in severe climate impacts, including increased "
     "extreme weather events, sea-level rise, and loss of biodiversity."),
]

# ── Susceptibility colour map (shared by appendix tables) ─────────────────────
_SUSC_BG: dict[str, str] = {
    "High Susceptibility":     "FCE4D6",
    "Moderate Susceptibility": "FFF2CC",
    "Low Susceptibility":      "E2EFDA",
}


# ── Appendix helper tables ────────────────────────────────────────────────────

def _flood_param_table(doc: Document, stats: dict) -> None:
    """Combined DEM + TWI susceptibility table for Appendix A."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A

    cls_order = ["High Susceptibility", "Moderate Susceptibility", "Low Susceptibility"]
    tbl = doc.add_table(rows=len(cls_order) + 1, cols=3)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Inches(2.0), Inches(2.25), Inches(2.25)]
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_widths[ci]

    for ci, hdr in enumerate(["Susceptibility Class",
                               "DEM — Elevation Range (m)", "TWI Range"]):
        style_cell(tbl.rows[0].cells[ci], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    for ri, cls in enumerate(cls_order):
        row  = tbl.rows[ri + 1]
        bg   = _SUSC_BG.get(cls, "FFFFFF")
        style_cell(row.cells[0], cls, bold=True, size=9, bg_color=bg)
        for ci, key in enumerate(["dem", "twi"], start=1):
            s = stats.get(key, {}).get(cls, {})
            txt = f'{s.get("range","—")}\n({s.get("pct","—")}%)' if s else "—"
            style_cell(row.cells[ci], txt, size=9, align=_A.CENTER)
    doc.add_paragraph()


def _heat_param_table(doc: Document, stats: dict) -> None:
    """Combined NDVI + NDBI + LST susceptibility table for Appendix A."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A

    cls_order = ["High Susceptibility", "Moderate Susceptibility", "Low Susceptibility"]
    tbl = doc.add_table(rows=len(cls_order) + 1, cols=4)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Inches(1.8), Inches(1.6), Inches(1.6), Inches(1.5)]
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_widths[ci]

    for ci, hdr in enumerate(["Susceptibility Class",
                               "NDVI Range", "NDBI Range", "LST Range (°C)"]):
        style_cell(tbl.rows[0].cells[ci], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    for ri, cls in enumerate(cls_order):
        row = tbl.rows[ri + 1]
        bg  = _SUSC_BG.get(cls, "FFFFFF")
        style_cell(row.cells[0], cls, bold=True, size=9, bg_color=bg)
        for ci, key in enumerate(["ndvi", "ndbi", "lst"], start=1):
            s = stats.get(key, {}).get(cls, {})
            txt = f'{s.get("range","—")}\n({s.get("pct","—")}%)' if s else "—"
            style_cell(row.cells[ci], txt, size=9, align=_A.CENTER)
    doc.add_paragraph()


def _susc_table(doc: Document, susc: dict) -> None:
    """3-class susceptibility distribution table (range + coverage %)."""
    if not susc:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A

    tbl = doc.add_table(rows=len(susc) + 1, cols=3)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Inches(2.5), Inches(2.5), Inches(1.5)]
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_widths[ci]

    for ci, hdr in enumerate(["Susceptibility Class", "Value Range", "Coverage (%)"]):
        style_cell(tbl.rows[0].cells[ci], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    for ri, (cls, data) in enumerate(susc.items()):
        row = tbl.rows[ri + 1]
        bg  = _SUSC_BG.get(cls, "F0F0F0")
        style_cell(row.cells[0], cls,                        bold=True, size=9, bg_color=bg)
        style_cell(row.cells[1], str(data.get("range","—")), size=9,   align=_A.CENTER)
        style_cell(row.cells[2], f'{data.get("pct","—")}%',  size=9,   align=_A.CENTER)
    doc.add_paragraph()


def _lulc_table(doc: Document, susc: dict) -> None:
    """LULC class coverage table."""
    if not susc:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A

    tbl = doc.add_table(rows=len(susc) + 1, cols=2)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    col_widths = [Inches(3.5), Inches(3.0)]
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_widths[ci]

    for ci, hdr in enumerate(["LULC Class", "Coverage (%)"]):
        style_cell(tbl.rows[0].cells[ci], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    for ri, (label, data) in enumerate(susc.items()):
        row = tbl.rows[ri + 1]
        style_cell(row.cells[0], label,                      size=9)
        style_cell(row.cells[1], f'{data.get("pct","—")}%',  size=9, align=_A.CENTER)
    doc.add_paragraph()


def _vector_type_table(doc: Document, feature_stats: dict) -> None:
    """Feature type distribution table for vector layers (roads / waterways)."""
    by_type = feature_stats.get("by_type", {})
    if not by_type:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH as _A

    has_length = any("length_km" in v for v in by_type.values())
    n_cols = 4 if has_length else 3
    headers = (["Feature Type", "Count", "Coverage (%)", "Length (km)"]
               if has_length else ["Feature Type", "Count", "Coverage (%)"])
    col_widths = ([Inches(2.0), Inches(1.5), Inches(1.5), Inches(1.5)]
                  if has_length else [Inches(2.5), Inches(1.5), Inches(2.5)])

    tbl = doc.add_table(rows=len(by_type) + 1, cols=n_cols)
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_widths[ci]

    for ci, hdr in enumerate(headers):
        style_cell(tbl.rows[0].cells[ci], hdr, bold=True, size=9,
                   bg_color=HDR_BG_HEX, font_color=COL_WHITE)

    for ri, (ftype, data) in enumerate(by_type.items()):
        row = tbl.rows[ri + 1]
        style_cell(row.cells[0], str(ftype).capitalize(), bold=True, size=9)
        style_cell(row.cells[1], str(data.get("count", "—")),    size=9, align=_A.CENTER)
        style_cell(row.cells[2], f'{data.get("pct", "—")}%',     size=9, align=_A.CENTER)
        if has_length:
            lkm = data.get("length_km")
            style_cell(row.cells[3],
                       f'{lkm:.2f}' if isinstance(lkm, (int, float)) else "—",
                       size=9, align=_A.CENTER)
    doc.add_paragraph()


def _appendix_layer_block(
    doc: Document,
    key: str,
    name: str,
    fig_ref: str,
    map_paths: dict,
    stats: dict,
    impact_text: str = "",
    categorical: bool = False,
    vector: bool = False,
) -> None:
    """Render one layer block: subheading → description → map → impact → table."""
    doc.add_paragraph()
    add_subheading(doc, f"Fig {fig_ref} — {name}", size=10, color=COL_HEADING)

    txt = _LAYER_APPENDIX_TEXT.get(key, {})
    if txt.get("description"):
        add_body(doc, txt["description"])

    png = map_paths.get(key)
    if png and Path(png).exists():
        add_image(doc, png, width=Inches(5.5),
                  caption=f"Figure {fig_ref} — {name}")

    # Use LLM-generated impact if available, fall back to static text
    final_impact = impact_text or txt.get("impact", "")
    if final_impact:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_run(p, "Impact: ", bold=True, size=10)
        add_run(p, final_impact, size=10)

    layer_stats = stats.get(key, {})
    if vector:
        _vector_type_table(doc, layer_stats)
    elif categorical:
        _lulc_table(doc, layer_stats)
    else:
        _susc_table(doc, layer_stats)


# ── Section 7 assembler ───────────────────────────────────────────────────────

def _appendices_section(doc: Document, ctx: dict) -> None:
    """Section 7 — Appendices."""
    map_paths     = ctx.get("appendix_map_paths",    {})
    stats         = ctx.get("appendix_stats",        {})
    layer_impacts = ctx.get("appendix_layer_impacts", {})
    risk_for      = ctx["risk_for"]

    if not map_paths and not stats:
        return

    add_page_break(doc)
    add_heading(doc, "7.  Appendices", size=13)

    # ── Appendix A ────────────────────────────────────────────────────────────
    add_subheading(doc, "Appendix A: Risk Assessment Data", size=12, color=COL_HEADING)
    add_body(doc,
        "This appendix presents the geospatial layers used to compute climate risk scores "
        "in this report. Each parameter is visualised as a spatial map alongside its "
        "susceptibility distribution, providing full transparency into the data that "
        "drives the ResSolv™ risk quantification methodology."
    )

    # Combined parameter summary tables
    if risk_for in ("Flood", "Both") and any(k in stats for k in ("dem", "twi")):
        add_subheading(doc, "Flood — Parameter Susceptibility Summary",
                       size=10, color=_FLOOD_ACCENT)
        _flood_param_table(doc, stats)

    if risk_for in ("Heat", "Both") and any(k in stats for k in ("ndvi", "ndbi", "lst")):
        add_subheading(doc, "Heat — Parameter Susceptibility Summary",
                       size=10, color=_HEAT_ACCENT)
        _heat_param_table(doc, stats)

    # Per-layer detailed blocks
    if risk_for in ("Flood", "Both"):
        for key, fig_ref, name in [
            ("dem",        "A.1(a)", "Elevation (DEM)"),
            ("twi",        "A.1(b)", "Topographic Wetness Index (TWI)"),
            ("impervious", "A.1(c)", "Impervious Surface Cover"),
        ]:
            _appendix_layer_block(doc, key, name, fig_ref, map_paths, stats,
                                  impact_text=layer_impacts.get(key, ""))

    if risk_for in ("Heat", "Both"):
        for key, fig_ref, name, cat in [
            ("ndvi", "A.2(a)", "Normalised Difference Vegetation Index (NDVI)", False),
            ("ndbi", "A.2(b)", "Normalised Difference Built-up Index (NDBI)",   False),
            ("lst",  "A.3",    "Land Surface Temperature (LST)",                False),
            ("lulc", "A.4",    "Land Use / Land Cover (LULC)",                  True),
        ]:
            _appendix_layer_block(doc, key, name, fig_ref, map_paths, stats,
                                  impact_text=layer_impacts.get(key, ""),
                                  categorical=cat)

    # ── Infrastructure layers (always shown when data is available) ───────────
    for key, fig_ref, name in [
        ("roads",     "A.5", "Road Network"),
        ("waterline", "A.6", "Waterways & Water Bodies"),
    ]:
        if key in map_paths or key in stats:
            _appendix_layer_block(doc, key, name, fig_ref, map_paths, stats,
                                  impact_text=layer_impacts.get(key, ""),
                                  vector=True)

    # ── Appendix A conclusion ─────────────────────────────────────────────────
    p_conc = doc.add_paragraph()
    p_conc.paragraph_format.space_after = Pt(8)
    add_run(p_conc, "Conclusion", bold=True, size=11,
            color=RGBColor(0x71, 0x56, 0xE6))
    hazard_types = _HAZARD_LABEL.get(risk_for, risk_for.lower())
    add_body(doc, _APPENDIX_CONCLUSION_TEMPLATE.format(hazard_types=hazard_types))

    # ── Appendix B ────────────────────────────────────────────────────────────
    add_page_break(doc)
    add_subheading(doc, "Appendix B: Shared Socioeconomic Pathways", size=12, color=COL_HEADING)
    for entry_type, text in _SSP_APPENDIX_CONTENT:
        if entry_type == "body":
            add_body(doc, text)
        elif entry_type == "ssp_heading":
            add_subheading(doc, text, size=10, color=COL_HEADING)
        elif entry_type == "bullet":
            add_bullet(doc, text, size=10)


# ═════════════════════════════════════════════════════════════════════════════
# PUBLIC API
# ═════════════════════════════════════════════════════════════════════════════

def build(context: dict) -> Path | None:
    """
    Assemble and save the final Word report.

    Returns:
        Path to saved .docx, or None if GENERATE_WORD is False.
    """
    if not config.GENERATE_WORD:
        logger.info("GENERATE_WORD=False — Word document skipped.")
        return None

    doc = Document()
    set_doc_margins(doc)
    set_default_font(doc)
    add_header_footer(doc, "The document is confidential and proprietary to Resilience AI")

    logger.info("Assembling Word document...")
    _cover_page(doc, context)
    _intro_methodology(doc, context)
    _risk_overview(doc, context)
    _risk_summary_table(doc, context)
    _ssp_section(doc, context)
    _impact_section(doc, context)
    _historical_section(doc, context)
    _influencing_factors_section(doc, context)
    _appendices_section(doc, context)

    out_path = context["output_dir"] / "final_report.docx"
    doc.save(str(out_path))
    logger.info("Word report saved: %s", out_path)
    return out_path
